import os
from docx import Document
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache

# 从环境变量获取API密钥

API_KEY = "xx"

# 使用指定的API URL和模型
client = OpenAI(api_key=API_KEY, base_url="https://api.deepseek.com")

@lru_cache(maxsize=None)
def translate_text(text):
    print(f"正在翻译文本: {text[:50]}...")  # 提示正在翻译的文本片段
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": f"请将以下内容翻译为另一种语言，保持格式不变:\n{text}"},
        ],
        temperature=0.7
    )
    translated_text = response.choices[0].message.content.strip()
    print(f"翻译完成: {translated_text[:50]}...")  # 提示翻译完成的文本片段
    return translated_text

def split_text(text, max_length=2000):
    texts = []
    current = []
    current_len = 0
    for line in text.split('\n'):
        if current_len + len(line) > max_length:
            texts.append('\n'.join(current))
            current = [line]
            current_len = len(line)
        else:
            current.append(line)
            current_len += len(line)
    if current:
        texts.append('\n'.join(current))
    return texts

def process_paragraph(paragraph):
    full_text = paragraph.text
    translated_parts = [translate_text(part) for part in split_text(full_text)]
    translated_text = "\n".join(translated_parts)
    paragraph.text = translated_text

def process_run(run):
    run_text = run.text
    translated_parts = [translate_text(part) for part in split_text(run_text)]
    translated_text = "\n".join(translated_parts)
    run.text = translated_text

def process_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph)

def process_document(input_file, output_file):
    # 检查输入文件
    if not os.path.isfile(input_file):
        print(f"错误：文件 {input_file} 不存在")
        return
    
    print(f"开始处理文件: {input_file}")  # 提示开始处理文件
    
    # 创建输出目录
    output_dir = os.path.dirname(os.path.abspath(output_file))
    os.makedirs(output_dir, exist_ok=True)
    
    # 加载源文件和创建目标文件
    doc_source = Document(input_file)
    doc_target = Document()  # 创建空目标文档
    
    # 并发处理每个段落和表格
    with ThreadPoolExecutor(max_workers=4) as executor:
        # 处理段落
        futures_paragraphs = [
            executor.submit(process_paragraph, paragraph)
            for paragraph in doc_source.paragraphs
        ]
        
        # 处理表格
        futures_tables = [
            executor.submit(process_table, table)
            for table in doc_source.tables
        ]
        
        for future in futures_paragraphs + futures_tables:
            future.result()  # 等待所有任务完成
    
    # 将处理后的内容添加到目标文档
    for paragraph in doc_source.paragraphs:
        new_paragraph = doc_target.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
    
    for table in doc_source.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        new_table = doc_target.add_table(rows=rows, cols=cols)
        for i_row, row in enumerate(table.rows):
            for i_col, cell in enumerate(row.cells):
                new_cell = new_table.cell(i_row, i_col)
                for paragraph in cell.paragraphs:
                    new_paragraph = new_cell.add_paragraph(paragraph.text)
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
    
    # 保存结果
    doc_target.save(output_file)
    print(f"翻译完成，结果已保存到 {output_file}")  # 提示翻译完成并保存文件

if __name__ == "__main__":
    INPUT_FILE = input("请输入要翻译的Word文件名: ")
    OUTPUT_FILE = INPUT_FILE.rsplit('.', 1)[0] + "_fy.docx"
    try:
        process_document(INPUT_FILE, OUTPUT_FILE)
    except Exception as e:
        print(f"发生错误: {e}")



